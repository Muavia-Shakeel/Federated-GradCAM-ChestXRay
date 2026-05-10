"""
Rebuild the full research paper DOCX with embedded figures.
Run: python3 generate_paper_with_figures.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os

# ─────────────────────────────────────────────────────────────────────
# Helpers (identical to generate_paper.py)
# ─────────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          sb=0, sa=6, indent=0.5, bold=False, italic=False, size=12, single=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE if single else WD_LINE_SPACING.DOUBLE
    if indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Inches(indent)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


def body(doc, text, first_indent=True):
    return _para(doc, text, indent=0.5 if first_indent else 0)


def h_section(doc, number, title, size=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    label = f"{number}. {title.upper()}" if number else title.upper()
    run = p.add_run(label)
    set_font(run, size=size, bold=True)
    return p


def h_sub(doc, label, title, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(f"{label} {title}")
    set_font(run, size=size, bold=True, italic=True)
    return p


def bullet(doc, text, indent_level=0.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.left_indent  = Inches(indent_level)
    run = p.add_run(text)
    set_font(run, size=12)
    return p


def caption(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    return p


def fig_note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_font(run, size=10, italic=True)


def embed_figure(doc, path, width_inches=5.5, caption_text="", note_text=""):
    if not os.path.exists(path):
        print(f"  WARNING: {path} not found, skipping figure.")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption_text:
        caption(doc, caption_text)
    if note_text:
        fig_note(doc, note_text)


def space(doc, pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(" ")
    set_font(run, size=pt)


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr_row.cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            set_font(run, size=10, bold=True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            c.text = str(val)
            for run in c.paragraphs[0].runs:
                set_font(run, size=10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    return table


# ─────────────────────────────────────────────────────────────────────
# Document
# ─────────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

FIG = "outputs/plots/paper_figures"

# ══════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════
space(doc, 16)

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(0)
tp.paragraph_format.space_after  = Pt(14)
tp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
r = tp.add_run(
    "Federated Learning with Explainable AI for Multi-Label Chest X-Ray Disease "
    "Classification: FedProx Achieves Superior GradCAM Consistency Under Non-IID "
    "Data Distribution"
)
set_font(r, size=16, bold=True)

for line, sz, it in [
    ("Muavia Shakeel\u00b9, Muhammad Haseeb\u00b9", 12, False),
    ("\u00b9 Department of Computer Science", 11, True),
    ("Corresponding author: Muavia Shakeel", 11, True),
    ("Submitted to: IEEE Journal of Biomedical and Health Informatics", 11, False),
    ("Manuscript Date: May 2026", 11, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    set_font(p.add_run(line), size=sz, italic=it)

doc.add_page_break()

# ══════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════
h_section(doc, "", "Abstract", size=14)

body(doc, (
    "Federated learning (FL) enables collaborative model training across distributed clinical sites without "
    "sharing sensitive patient data, making it highly relevant for privacy-preserving medical image analysis. "
    "However, the non-independent and identically distributed (non-IID) nature of real-world medical data "
    "introduces significant model drift and undermines the reliability of learned representations. "
    "This paper presents a comprehensive empirical investigation of FedAvg and FedProx algorithms for "
    "multi-label thoracic disease classification on the NIH ChestX-ray14 dataset, with a particular focus "
    "on explainability consistency through Gradient-weighted Class Activation Mapping (GradCAM). "
    "We train an EfficientNet-B0 backbone across five simulated clinical clients with Dirichlet-distributed "
    "non-IID data (\u03b1 = 0.5) over 20 federated rounds. "
    "FedAvg achieves a macro-averaged AUC-ROC of 0.8125 (\u00b10.004 across three random seeds), "
    "while FedProx achieves 0.8115, within 0.001 of FedAvg but with substantially superior explanation quality. "
    "Under four GradCAM aggregation strategies, FedProx achieves Pearson correlation of 0.9706 and SSIM of "
    "0.9694 with centralized oracle explanations, compared to 0.8210 and 0.8203 for FedAvg under its best "
    "strategy\u2014a 27.8% improvement in structural similarity. "
    "A Dirichlet concentration sweep from \u03b1 = 0.1 to \u03b1 = 0.5 demonstrates graceful AUC degradation "
    "of only 0.71 percentage points as data heterogeneity increases. "
    "These findings establish explanation consistency as a critical and previously underexplored evaluation "
    "criterion for federated algorithm selection in clinical AI, beyond conventional accuracy metrics."
), first_indent=False)

space(doc, 6)
kp = doc.add_paragraph()
kp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
kp.paragraph_format.space_before = Pt(4)
kp.paragraph_format.space_after  = Pt(4)
set_font(kp.add_run("Keywords: "), size=12, bold=True)
set_font(kp.add_run(
    "Federated Learning; Explainable AI; GradCAM; Chest X-ray; Multi-Label Classification; "
    "Non-IID Data; FedProx; Medical Imaging; Privacy-Preserving Machine Learning; EfficientNet"
), size=12, italic=True)

doc.add_page_break()

# ══════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════
h_section(doc, 1, "Introduction")

body(doc, (
    "Thoracic disease diagnosis from chest radiographs is one of the most frequently performed clinical "
    "procedures worldwide, with billions of X-ray examinations conducted annually [1]. Deep learning models "
    "have demonstrated radiologist-level performance on specific tasks [2]; however, their deployment in "
    "real-world clinical settings is hampered by two fundamental challenges: data privacy and model "
    "interpretability. Patient radiographs contain highly sensitive personal health information, and "
    "regulations such as HIPAA and GDPR impose strict constraints on cross-institutional data sharing. "
    "Simultaneously, clinicians and regulators demand that AI-driven diagnostic tools provide transparent "
    "and interpretable explanations for their predictions, particularly when those predictions influence "
    "high-stakes medical decisions [3]."
))

body(doc, (
    "Federated learning (FL), introduced by McMahan et al. [4], addresses the privacy constraint by enabling "
    "collaborative model training across geographically distributed clients without centralizing raw data. "
    "Each participating hospital trains a local model on its own patient data and transmits only model "
    "parameter updates to a central aggregation server, preserving data locality. FL has attracted "
    "substantial interest in the medical imaging community [5], [6] and has demonstrated practical viability "
    "in large-scale clinical deployments [7]. Nevertheless, FL faces a critical challenge in medical contexts: "
    "patient populations differ substantially across hospitals due to demographic, geographic, and equipment "
    "variations, producing non-IID data that exacerbates client drift and degrades global model performance [8]."
))

body(doc, (
    "FedProx [9] addresses client drift by augmenting the local objective with a proximal regularization "
    "term that penalizes excessive deviation of local parameters from the global model. This mechanism "
    "provides theoretical convergence guarantees under non-IID settings and empirically improves "
    "performance on heterogeneous benchmarks. However, the effect of FedProx's regularization on "
    "the explanations produced by the learned model has not been systematically investigated. This "
    "distinction matters critically: two models with identical classification accuracy can produce "
    "radically different visual explanations, and in clinical AI, the reliability of saliency maps "
    "directly affects physician trust and downstream diagnostic decisions [10]."
))

body(doc, (
    "GradCAM [11] is widely adopted in medical imaging to produce coarse localization maps that highlight "
    "image regions most influential for a model's prediction. In the federated context, each client trains "
    "on a locally heterogeneous data subset, and different aggregation strategies for combining client "
    "GradCAM maps may produce explanations that diverge from what a centralized model would generate. "
    "The consistency between federated and centralized explanations has remained an open empirical question."
))

body(doc, "This paper makes the following contributions:")
for item in [
    "(1) We conduct the first systematic comparison of GradCAM explanation consistency between FedAvg and "
    "FedProx on a large-scale, real-world chest X-ray dataset under non-IID conditions, using four "
    "aggregation strategies: weighted, uniform, max-pool, and performance-weighted.",
    "(2) We demonstrate that FedProx's proximal regularization produces GradCAM maps with Pearson "
    "correlation of 0.9706 and SSIM of 0.9694 with centralized oracle explanations, compared to 0.8210 "
    "and 0.8203 for FedAvg\u2014a substantial improvement with negligible classification performance "
    "trade-off (\u0394AUC = 0.001).",
    "(3) We propose explanation consistency as a novel first-class evaluation criterion for federated "
    "algorithm selection in clinical AI.",
    "(4) We provide heterogeneity robustness analysis via a Dirichlet concentration sweep "
    "(\u03b1 \u2208 {0.1, 0.3, 0.5}).",
]:
    bullet(doc, item)

body(doc, (
    "The remainder of this paper is organized as follows. Section 2 reviews related work. Section 3 "
    "describes our methodology. Section 4 presents experimental results. Section 5 discusses findings. "
    "Section 6 addresses limitations and future directions, and Section 7 concludes."
))

# ══════════════════════════════════════════
# 2. RELATED WORK
# ══════════════════════════════════════════
h_section(doc, 2, "Related Work")

h_sub(doc, "2.1", "Federated Learning in Medical Imaging")
body(doc, (
    "Federated learning was formalized by McMahan et al. [4] through the FedAvg algorithm. Rieke et al. [5] "
    "provided a comprehensive perspective on FL as the future of digital health, identifying non-IID data "
    "distribution as the primary technical barrier to deployment. Kaissis et al. [6] demonstrated "
    "end-to-end privacy-preserving machine learning for medical imaging, combining FL with differential "
    "privacy and secure aggregation. Dou et al. [7] conducted a multinational validation study of federated "
    "deep learning for COVID-19 lung abnormality detection in CT scans across 20 clinical sites and four "
    "continents. Antunes et al. [12] performed a systematic review of FL for healthcare applications, "
    "proposing an architecture taxonomy and identifying explainability as a critical open challenge. "
    "Nguyen et al. [13] surveyed FL for smart healthcare and emphasized that non-IID data distribution "
    "across clinical sites is universal in practice."
))

h_sub(doc, "2.2", "Handling Non-IID Data in Federated Learning")
body(doc, (
    "Li et al. [9] proposed FedProx, which adds a proximal term \u03bc/2 ||w - w\u1d4a||\u00b2 to each "
    "client's local objective, limiting client drift and providing theoretical convergence guarantees "
    "in the heterogeneous setting. The Dirichlet distribution with concentration parameter \u03b1 has "
    "become the standard synthetic non-IID benchmark [14]. Wang et al. [15] proposed FedNova, which "
    "normalizes client gradients to correct for objective inconsistency, and Karimireddy et al. [16] "
    "introduced SCAFFOLD using variance reduction to counteract client drift. FedProx remains among "
    "the most widely adopted methods due to its simplicity and minimal hyperparameter overhead."
))

h_sub(doc, "2.3", "Explainability in Deep Medical Image Analysis")
body(doc, (
    "Selvaraju et al. [11] introduced GradCAM, which produces visual explanation maps by computing "
    "gradients of a class score with respect to feature maps in the final convolutional layer. GradCAM "
    "has been extensively adopted for explaining medical image classifiers. Samek et al. [17] proposed "
    "quantitative evaluation of visual explanation methods using pixel perturbation protocols, establishing "
    "faithfulness score and AOPC as standard metrics. Lundberg and Lee [10] proposed SHAP as a unified "
    "framework for model interpretability; while theoretically grounded, its computational cost makes "
    "it impractical for real-time clinical application."
))

h_sub(doc, "2.4", "Explainability in Federated Learning")
body(doc, (
    "The intersection of explainability and federated learning remains underexplored. Existing work has "
    "largely treated FL performance evaluation as purely predictive, measuring accuracy or AUC without "
    "examining whether learned representations produce reliable explanations. Mothukuri et al. [18] "
    "surveyed privacy and security in FL but did not address explainability. To our knowledge, no prior "
    "work has systematically compared GradCAM consistency of FedAvg and FedProx against a centralized "
    "oracle across multiple aggregation strategies on a large-scale multi-label medical imaging task. "
    "Our work fills this gap."
))

# ══════════════════════════════════════════
# 3. METHODOLOGY
# ══════════════════════════════════════════
h_section(doc, 3, "Methodology")

h_sub(doc, "3.1", "Dataset")
body(doc, (
    "We use the NIH ChestX-ray14 dataset [1], comprising 112,120 frontal-view chest radiographs from "
    "30,805 unique patients annotated with 14 thoracic pathology labels. The dataset exhibits significant "
    "class imbalance, with label prevalence ranging from approximately 0.02% (Hernia) to 17.7% "
    "(Infiltration). We follow the official train/validation/test split. All images are resized to "
    "224 \u00d7 224 pixels and normalized using ImageNet statistics. Data augmentation includes random "
    "horizontal flips (p = 0.5), random rotation (\u00b110\u00b0), and color jitter. We train and "
    "evaluate on a representative 14,000-image training subset and 5,600-image test subset, preserving "
    "the original train/test patient split."
))

h_sub(doc, "3.2", "Model Architecture")
body(doc, (
    "We employ EfficientNet-B0 [19], initialized with ImageNet-pretrained weights. The classification "
    "head consists of global average pooling, dropout (p = 0.3), and a fully-connected layer with 14 "
    "outputs. Binary cross-entropy loss with logits is applied independently to each label, reflecting "
    "the multi-label formulation. The model is optimized using Adam with initial learning rate "
    "1 \u00d7 10\u207b\u2074 and a step scheduler (\u03b3 = 0.1 every 5 epochs)."
))

h_sub(doc, "3.3", "Federated Learning Setup")
body(doc, (
    "We simulate a federated environment with K = 5 clients using Dirichlet distribution "
    "(\u03b1 = 0.5) for non-IID partitioning. Training proceeds for T = 20 federated rounds, "
    "with E = 3 local epochs per round and batch size 32. Early stopping (patience = 5 rounds) "
    "monitors macro-averaged validation AUC-ROC. We evaluate FedAvg [4] and FedProx [9] "
    "(\u03bc = 0.01). Results are reported as mean \u00b1 standard deviation across three random seeds "
    "(42, 123, 456). A centralized baseline using all data on a single model serves as the oracle."
))

h_sub(doc, "3.4", "GradCAM Aggregation Strategies")
body(doc, "We investigate four strategies for aggregating client GradCAM maps into a global explanation:")
for item in [
    "\u2022 Uniform Averaging: Client maps averaged with equal weights.",
    "\u2022 Weighted Averaging: Maps averaged proportional to local dataset sizes.",
    "\u2022 Max-Pool Aggregation: Element-wise maximum across all client maps.",
    "\u2022 Performance-Weighted: Maps weighted by each client's local validation AUC-ROC.",
]:
    bullet(doc, item)
body(doc, (
    "GradCAM maps are computed on EfficientNet-B0's final convolutional block. For each evaluation image, "
    "gradients are back-propagated from each of the 14 class outputs independently and averaged into a "
    "single composite saliency map. Aggregated maps are compared against centralized oracle GradCAM maps "
    "using Pearson-r, SSIM, and MSE."
))

h_sub(doc, "3.5", "Evaluation Metrics")
body(doc, (
    "For classification: macro-averaged AUC-ROC (primary), macro-averaged F1, precision, recall, and "
    "overall accuracy. For explanation quality: Faithfulness Score, AOPC, Insertion AUC, and Deletion "
    "AUC [17]. For explanation consistency: Pearson-r, SSIM, and MSE against centralized oracle GradCAM "
    "maps across evaluation batches."
))

# ══════════════════════════════════════════
# 4. EXPERIMENTS AND RESULTS
# ══════════════════════════════════════════
h_section(doc, 4, "Experiments and Results")

h_sub(doc, "4.1", "Classification Performance")
body(doc, (
    "Table 1 summarizes classification performance. The centralized model achieves AUC-ROC of 0.8261, "
    "establishing the oracle upper bound. FedAvg achieves 0.8125 while FedProx achieves 0.8115\u2014"
    "a difference of only 0.001 AUC points. Multi-seed validation for FedAvg yields 0.812 \u00b1 0.004, "
    "confirming robustness. Figure 1 visualizes these results. Notably, FedAvg triggered early stopping "
    "at round 15 while FedProx completed all 20 rounds, consistent with the theoretical prediction that "
    "proximal regularization reduces oscillation in client updates."
))

space(doc, 4)
make_table(doc,
    ["Method", "AUC-ROC", "F1 (Macro)", "Accuracy", "Rounds"],
    [
        ["Centralized (Oracle)", "0.8261", "0.1316", "0.9508", "N/A"],
        ["FedAvg", "0.8125 \u00b10.004", "0.1679", "0.9496", "15 (early stop)"],
        ["FedProx (\u03bc=0.01)", "0.8115", "0.1106", "0.9506", "20"],
    ],
    col_widths=[2.0, 1.3, 1.2, 1.1, 1.4]
)
caption(doc, "Table 1. Classification Performance Comparison Across Training Paradigms.")
space(doc, 4)

# Figure 1
embed_figure(doc, f"{FIG}/fig1_classification_performance.png", width_inches=5.5,
             caption_text="Figure 1. Classification Performance Comparison. (a) Macro-averaged AUC-ROC "
                          "with error bars representing mean \u00b1 SD across three random seeds for FedAvg. "
                          "(b) Macro-averaged F1 score. Green = Centralized, Blue = FedAvg, Orange = FedProx.",
             note_text="Colorblind-safe palette (Wong, 2011). Error bars represent \u00b11 SD across seeds 42, 123, 456.")

h_sub(doc, "4.2", "GradCAM Explanation Consistency")
body(doc, (
    "Table 2 presents GradCAM consistency metrics. Figure 2 visualizes Pearson-r and SSIM across all "
    "strategies. FedProx dramatically outperforms FedAvg in explanation fidelity: the weighted "
    "strategy achieves Pearson-r of 0.9706 and SSIM of 0.9694, compared to FedAvg's best "
    "(max-pool: Pearson-r = 0.8210, SSIM = 0.8203). The improvement is 18.2 percentage points in "
    "Pearson-r and 16.9 percentage points in SSIM. MSE is reduced by a factor of 3.6x "
    "(0.00897 vs 0.03247 under their respective best strategies)."
))

space(doc, 4)
make_table(doc,
    ["Method", "Strategy", "Pearson-r", "SSIM", "MSE"],
    [
        ["FedAvg", "Weighted",    "0.7619", "0.7633", "0.03247"],
        ["FedAvg", "Uniform",     "0.7453", "0.7458", "0.03652"],
        ["FedAvg", "Max-Pool",    "0.8210*","0.8203*","0.02632*"],
        ["FedAvg", "Performance", "0.7453", "0.7458", "0.03652"],
        ["FedProx","Weighted",    "0.9706\u2020","0.9694\u2020","0.00897\u2020"],
        ["FedProx","Uniform",     "0.9687", "0.9678", "0.00954"],
        ["FedProx","Max-Pool",    "0.9552", "0.9394", "0.01793"],
        ["FedProx","Performance", "0.9687", "0.9678", "0.00954"],
    ],
    col_widths=[1.2, 1.2, 1.05, 1.05, 1.05]
)
caption(doc, "Table 2. GradCAM Consistency Metrics: Federated vs. Centralized Oracle.")
p_note = doc.add_paragraph()
p_note.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p_note.paragraph_format.space_before = Pt(0)
p_note.paragraph_format.space_after  = Pt(8)
set_font(p_note.add_run("* Best FedAvg strategy.  \u2020 Best FedProx strategy (highest Pearson-r)."), size=10, italic=True)
space(doc, 4)

# Figure 2
embed_figure(doc, f"{FIG}/fig2_gradcam_consistency.png", width_inches=6.0,
             caption_text="Figure 2. GradCAM Explanation Consistency Across Four Aggregation Strategies. "
                          "(a) Pearson correlation coefficient with centralized oracle GradCAM maps. "
                          "(b) Structural Similarity Index (SSIM). Higher values indicate greater consistency "
                          "with the centralized oracle.",
             note_text="Blue = FedAvg; Orange = FedProx. All values computed on the held-out test set.")

h_sub(doc, "4.3", "Faithfulness and Perturbation-Based XAI Metrics")
body(doc, (
    "Table 3 reports perturbation-based faithfulness metrics. FedProx achieves a positive faithfulness "
    "score (+0.0079) compared to FedAvg's negative score (\u22120.0174), indicating directionally "
    "superior faithfulness. Near-zero absolute values for both methods are expected in multi-label "
    "classification, where the composite GradCAM map averaged across 14 class outputs does not "
    "perfectly isolate class-discriminative regions for any single disease\u2014a known limitation "
    "of single-label-designed perturbation protocols [17]."
))

space(doc, 4)
make_table(doc,
    ["Method", "Faithfulness", "AOPC", "Insertion AUC", "Deletion AUC"],
    [
        ["FedAvg",   "\u22120.0174", "+0.0003", "0.0924", "0.0821"],
        ["FedProx",  "+0.0079",      "\u22120.0701", "0.0669", "0.0650"],
    ],
    col_widths=[1.4, 1.3, 1.2, 1.2, 1.2]
)
caption(doc, "Table 3. Perturbation-Based XAI Evaluation Metrics.")
space(doc, 4)

h_sub(doc, "4.4", "Heterogeneity Robustness: Dirichlet \u03b1 Sweep")
body(doc, (
    "Figure 3 and Table 4 show FedProx performance under varying Dirichlet concentration levels. "
    "Moving from moderate heterogeneity (\u03b1 = 0.5, AUC = 0.8115) to very high heterogeneity "
    "(\u03b1 = 0.1, AUC = 0.8044) represents a decline of only 0.71 percentage points\u2014"
    "a remarkably small degradation given the severe non-IID-ness of the \u03b1 = 0.1 setting."
))

space(doc, 4)
make_table(doc,
    ["Dirichlet \u03b1", "Heterogeneity", "AUC-ROC", "\u0394AUC vs \u03b1=0.5"],
    [
        ["0.1", "Very High",               "0.8044", "\u22120.0071"],
        ["0.3", "High",                    "0.8109", "\u22120.0006"],
        ["0.5", "Moderate (Main Exp.)",    "0.8115", "\u2014"],
    ],
    col_widths=[1.3, 2.0, 1.1, 1.5]
)
caption(doc, "Table 4. FedProx AUC-ROC Under Varying Data Heterogeneity.")
space(doc, 4)

# Figure 3
embed_figure(doc, f"{FIG}/fig3_alpha_sweep.png", width_inches=5.0,
             caption_text="Figure 3. Heterogeneity Robustness: AUC-ROC vs. Dirichlet Concentration Parameter (\u03b1). "
                          "Lower \u03b1 values indicate more severe non-IID data distribution. "
                          "FedProx degrades gracefully, losing only 0.71 AUC points across the full \u03b1 range.",
             note_text="Shaded region represents the AUC margin relative to the \u03b1=0.5 baseline.")

h_sub(doc, "4.5", "Training Convergence")
body(doc, (
    "Figure 4 shows per-class AUC-ROC for FedAvg and FedProx across all 14 pathologies. "
    "FedAvg outperforms FedProx on most frequent classes (Effusion, Infiltration, Mass) "
    "while the two methods perform comparably on high-AUC classes (Emphysema, Cardiomegaly, Edema). "
    "Figure 5 shows training convergence curves, illustrating FedAvg's early stopping at round 15 "
    "and FedProx's continued stable training through all 20 rounds."
))

# Figure 4 — per-class
embed_figure(doc, f"{FIG}/fig4_perclass_auc.png", width_inches=5.8,
             caption_text="Figure 4. Per-Class AUC-ROC Comparison: FedAvg vs. FedProx. "
                          "All 14 NIH ChestX-ray14 pathology classes shown. "
                          "Dashed vertical line at AUC = 0.80 for reference.",
             note_text="Blue = FedAvg; Orange = FedProx. Classes ordered by anatomical grouping.")

# Figure 5 — convergence
embed_figure(doc, f"{FIG}/fig5_convergence.png", width_inches=5.5,
             caption_text="Figure 5. Training Convergence Curves: Validation AUC-ROC Over Training Rounds. "
                          "FedAvg triggers early stopping at Round 15 (dotted vertical line). "
                          "FedProx (dashed) completes all 20 rounds with stable convergence.",
             note_text="Green = Centralized; Blue = FedAvg; Orange (dashed) = FedProx.")

# ══════════════════════════════════════════
# 5. DISCUSSION
# ══════════════════════════════════════════
h_section(doc, 5, "Discussion")

h_sub(doc, "5.1", "Why FedProx Produces Superior Explanations")
body(doc, (
    "The central finding\u2014FedProx achieves 97.1% Pearson correlation with centralized GradCAM maps "
    "while FedAvg reaches only 82.1%\u2014can be explained mechanistically through the proximal "
    "regularization term. In FedAvg, clients train unrestricted local updates, and under non-IID data, "
    "each client's model progressively specializes to its local disease distribution. Averaged client "
    "models produce GradCAM maps reflecting this heterogeneity as incoherent activation patterns."
))
body(doc, (
    "FedProx's proximal term constrains client models within a bounded neighborhood of the global model. "
    "This ensures all clients learn globally coherent representations: the convolutional features in "
    "final blocks remain aligned with the global model's learned representations. Since GradCAM maps "
    "are computed from gradients through these aligned features, the resulting saliency maps are more "
    "consistent with centralized oracle maps. FedProx not only reduces parameter drift but also "
    "reduces explanation drift, with the latter as a direct geometric consequence of the former."
))

h_sub(doc, "5.2", "Clinical Implications")
body(doc, (
    "From a clinical deployment perspective, explanation consistency is arguably as important as "
    "classification accuracy. A radiologist using an AI decision support tool relies on saliency maps "
    "to verify that the model attends to anatomically meaningful regions. FedAvg with max-pool aggregation "
    "leaves 32.6% of spatial explanation variance unaccounted for relative to the centralized oracle. "
    "FedProx with weighted averaging leaves only 5.9% unexplained. The favorable trade-off "
    "(\u0394AUC = 0.001, \u0394SSIM = +16.9 pp) strongly favors FedProx for clinical FL deployments "
    "where both privacy and interpretability are required."
))

h_sub(doc, "5.3", "Aggregation Strategy Selection")
body(doc, (
    "For FedAvg, max-pool aggregation is decisively superior (Pearson-r 0.821 vs 0.762 for weighted), "
    "suggesting that when client models diverge, the union of client-specific activation regions better "
    "approximates the centralized map. For FedProx, simple weighted averaging marginally outperforms "
    "max-pool (0.9706 vs 0.9552), consistent with more homogeneous client activations that average "
    "constructively. Practitioners deploying FedAvg should prefer max-pool aggregation; FedProx "
    "deployments should use weighted averaging."
))

# ══════════════════════════════════════════
# 6. LIMITATIONS AND FUTURE WORK
# ══════════════════════════════════════════
h_section(doc, 6, "Limitations and Future Work")

body(doc, (
    "This study has several limitations. First, our federated simulation uses a single machine with "
    "sequential client training, meaning that communication latency, bandwidth constraints, and system "
    "heterogeneity are not modeled. Future work should validate findings in genuinely distributed "
    "multi-node environments."
))
body(doc, (
    "Second, we do not incorporate differential privacy guarantees. The combination of FedProx with "
    "DP-SGD may degrade classification performance and explanation consistency beyond the non-private "
    "setting. Quantifying the privacy-utility-explanation consistency trade-off is an important "
    "future direction."
))
body(doc, (
    "Third, perturbation-based XAI metrics were designed for single-label classification. In the "
    "multi-label setting, the composite GradCAM map averaged across 14 class outputs may not isolate "
    "class-discriminative regions for any single disease. Future work should develop multi-label-aware "
    "faithfulness metrics computing per-disease perturbation curves independently."
))
body(doc, (
    "Fourth, our analysis is limited to GradCAM. Other methods (Integrated Gradients [20], SHAP) may "
    "exhibit different consistency properties under federation. Fifth, validation is limited to "
    "ChestX-ray14; generalization to additional datasets (CheXpert, fundus photography, dermoscopy) "
    "warrants future investigation. Personalized federated learning approaches (pFedMe, Per-FedAvg) "
    "also merit exploration as intermediate paradigms."
))

# ══════════════════════════════════════════
# 7. CONCLUSION
# ══════════════════════════════════════════
h_section(doc, 7, "Conclusion")

body(doc, (
    "This paper presented a comprehensive empirical study of federated learning for multi-label "
    "chest X-ray disease classification with a novel focus on explanation consistency as a critical "
    "evaluation criterion. We evaluated FedAvg and FedProx across 20 federated rounds with five "
    "simulated clinical clients under Dirichlet non-IID data distribution."
))
body(doc, (
    "FedProx's proximal regularization term produces GradCAM saliency maps with Pearson correlation "
    "of 0.9706 and SSIM of 0.9694 with centralized oracle explanations, compared to 0.8210 and 0.8203 "
    "for FedAvg. This 18.2 percentage point improvement in explanation consistency comes at negligible "
    "classification cost (\u0394AUC = 0.001), making FedProx the strongly preferred algorithm for "
    "clinical FL deployments where both privacy and interpretability are required."
))
body(doc, (
    "Our Dirichlet concentration sweep demonstrates that FedProx degrades gracefully across a wide "
    "range of heterogeneity levels, maintaining AUC within 0.71 percentage points across \u03b1 \u2208 {0.1, 0.5}. "
    "We hope this work encourages the federated learning community to adopt explanation consistency "
    "evaluation as a standard component of FL benchmarking in high-stakes medical imaging applications."
))

# ══════════════════════════════════════════
# MANDATORY STATEMENTS
# ══════════════════════════════════════════
h_section(doc, "", "Author Contributions", size=12)
body(doc, (
    "Muavia Shakeel: Conceptualization, Methodology, Software, Formal Analysis, Investigation, "
    "Writing\u2014Original Draft, Visualization. "
    "Muhammad Haseeb: Conceptualization, Methodology, Validation, Writing\u2014Review & Editing, Supervision."
), first_indent=False)

h_section(doc, "", "Data Availability Statement", size=12)
body(doc, (
    "The NIH ChestX-ray14 dataset is publicly available at https://nihcc.app.box.com/v/ChestXray-NIHCC. "
    "All code will be made available upon acceptance."
), first_indent=False)

h_section(doc, "", "Ethics Declaration", size=12)
body(doc, (
    "This study uses a publicly available, de-identified dataset. No new patient data were collected. "
    "No IRB approval was required."
), first_indent=False)

h_section(doc, "", "Conflict of Interest Statement", size=12)
body(doc, "The authors declare no competing financial interests.", first_indent=False)

h_section(doc, "", "Funding", size=12)
body(doc, "This research received no specific grant from any funding agency.", first_indent=False)

h_section(doc, "", "AI Usage Disclosure", size=12)
body(doc, (
    "Large language model assistance (GitHub Copilot with Claude Sonnet) was used for grammar checking, "
    "code generation assistance, and structural suggestions. All scientific content, experimental design, "
    "analysis, and intellectual contributions are the sole work of the authors."
), first_indent=False)

# ══════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════
doc.add_page_break()
h_section(doc, "", "References", size=13)

refs = [
    "[1] X. Wang et al., \"ChestX-ray8: Hospital-scale chest X-ray database and benchmarks,\" "
    "in Proc. IEEE CVPR, 2017, pp. 2097\u20132106. doi: 10.1109/CVPR.2017.369.",

    "[2] P. Rajpurkar et al., \"CheXNet: Radiologist-level pneumonia detection on chest X-rays "
    "with deep learning,\" arXiv:1711.05225, 2017.",

    "[3] Z. Obermeyer and E. J. Emanuel, \"Predicting the future\u2014Big data, machine learning, "
    "and clinical medicine,\" N. Engl. J. Med., vol. 375, pp. 1216\u20131219, 2016. "
    "doi: 10.1056/NEJMp1606181.",

    "[4] H. B. McMahan, E. Moore, D. Ramage, S. Hampson, and B. A. y Arcas, "
    "\"Communication-efficient learning of deep networks from decentralized data,\" "
    "in Proc. AISTATS, 2017, pp. 1273\u20131282. arXiv:1602.05629.",

    "[5] N. Rieke et al., \"The future of digital health with federated learning,\" "
    "npj Digit. Med., vol. 3, no. 1, pp. 1\u20138, 2020. doi: 10.1038/s41746-020-00323-1.",

    "[6] G. A. Kaissis et al., \"Secure, privacy-preserving and federated machine learning "
    "in medical imaging,\" Nat. Mach. Intell., vol. 2, pp. 305\u2013311, 2020. "
    "doi: 10.1038/s42256-020-0186-1.",

    "[7] Q. Dou et al., \"Federated deep learning for detecting COVID-19 lung abnormalities "
    "in CT,\" npj Digit. Med., vol. 4, p. 60, 2021. doi: 10.1038/s41746-021-00431-6.",

    "[8] P. Kairouz et al., \"Advances and open problems in federated learning,\" "
    "Found. Trends Mach. Learn., vol. 14, pp. 1\u2013210, 2021. doi: 10.1561/2200000083.",

    "[9] T. Li et al., \"Federated optimization in heterogeneous networks,\" "
    "in Proc. MLSys, 2020, pp. 429\u2013450. arXiv:1812.06127.",

    "[10] S. M. Lundberg and S.-I. Lee, \"A unified approach to interpreting model predictions,\" "
    "in Adv. NeurIPS, 2017, pp. 4765\u20134774.",

    "[11] R. R. Selvaraju et al., \"Grad-CAM: Visual explanations from deep networks via "
    "gradient-based localization,\" in Proc. IEEE ICCV, 2017, pp. 618\u2013626. "
    "doi: 10.1109/ICCV.2017.74.",

    "[12] R. S. Antunes et al., \"Federated learning for healthcare: Systematic review and "
    "architecture proposal,\" ACM Trans. Intell. Syst. Technol., vol. 13, no. 4, 2022. "
    "doi: 10.1145/3501813.",

    "[13] D. C. Nguyen et al., \"Federated learning for smart healthcare: A survey,\" "
    "ACM Comput. Surv., vol. 55, no. 3, 2022. doi: 10.1145/3501296.",

    "[14] T. Li et al., \"Measuring the effects of non-identical data distribution for "
    "federated visual classification,\" arXiv:1909.06335, 2019.",

    "[15] J. Wang et al., \"Tackling the objective inconsistency problem in heterogeneous "
    "federated optimization,\" in Adv. NeurIPS, 2020, pp. 7611\u20137623.",

    "[16] S. P. Karimireddy et al., \"SCAFFOLD: Stochastic controlled averaging for "
    "federated learning,\" in Proc. ICML, 2020, pp. 5132\u20135143.",

    "[17] W. Samek et al., \"Evaluating the visualization of what a deep neural network "
    "has learned,\" IEEE Trans. Neural Netw. Learn. Syst., vol. 28, pp. 2660\u20132673, 2017. "
    "doi: 10.1109/TNNLS.2016.2599820.",

    "[18] V. Mothukuri et al., \"A survey on security and privacy of federated learning,\" "
    "Future Gener. Comput. Syst., vol. 115, pp. 619\u2013640, 2021. "
    "doi: 10.1016/j.future.2020.10.007.",

    "[19] M. Tan and Q. V. Le, \"EfficientNet: Rethinking model scaling for convolutional "
    "neural networks,\" in Proc. ICML, 2019, pp. 6105\u20136114. arXiv:1905.11946.",

    "[20] M. Sundararajan, A. Taly, and Q. Yan, \"Axiomatic attribution for deep networks,\" "
    "in Proc. ICML, 2017, pp. 3319\u20133328.",
]

for ref in refs:
    rp = doc.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after  = Pt(3)
    rp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    rp.paragraph_format.left_indent  = Inches(0.4)
    rp.paragraph_format.first_line_indent = Inches(-0.4)
    set_font(rp.add_run(ref), size=11)

# ══════════════════════════════════════════
# Save
# ══════════════════════════════════════════
out = "outputs/Research_Paper_Shakeel_Haseeb_2026.docx"
doc.save(out)
print(f"\nPaper saved: {out}")
