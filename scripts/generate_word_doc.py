"""
Convert Project_Complete_Document.md to a professional Word .docx file.
Run: /home/msi/DS_env/bin/python3 generate_word_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

DOC_PATH = os.path.join(os.path.dirname(__file__), "..", "docs", "Project_Complete_Document.md")
OUT_PATH = os.path.join(os.path.dirname(__file__), "..", "outputs", "reports", "Project_Complete_Document.docx")

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)

def set_heading_style(para, level, doc):
    style_name = f"Heading {level}"
    try:
        para.style = doc.styles[style_name]
    except Exception:
        pass
    run = para.runs[0] if para.runs else para.add_run()
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x8F)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x6E, 0x6E)

def add_horizontal_rule(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A1A6E')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table_from_md(doc, lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].startswith('|'):
        row = [c.strip() for c in lines[i].split('|')[1:-1]]
        rows.append(row)
        i += 1
    if not rows:
        return i
    # Filter separator row
    data_rows = [r for r in rows if not all(set(c) <= set('-: ') for c in r)]
    if not data_rows:
        return i
    header = data_rows[0]
    body = data_rows[1:]
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(header):
        hdr_cells[j].text = h
        run = hdr_cells[j].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        hdr_cells[j].paragraphs[0].paragraph_format.space_after = Pt(2)
        # Header background
        tc = hdr_cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E2F3')
        tcPr.append(shd)
    # Body rows
    for ri, row in enumerate(body):
        row_cells = table.rows[ri + 1].cells
        for j, cell_text in enumerate(row):
            if j < len(row_cells):
                row_cells[j].text = cell_text
                run = row_cells[j].paragraphs[0].runs[0] if row_cells[j].paragraphs[0].runs else row_cells[j].paragraphs[0].add_run(cell_text)
                run.font.size = Pt(9)
    doc.add_paragraph()
    return i

def parse_inline(text):
    """Return list of (text, bold, code) tuples."""
    segments = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], False, False))
        if m.group().startswith('**'):
            segments.append((m.group(2), True, False))
        else:
            segments.append((m.group(3), False, True))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False, False))
    return segments

def add_formatted_para(doc, text, style=None, font_size=11):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(4)
    segments = parse_inline(text)
    for seg_text, bold, code in segments:
        run = para.add_run(seg_text)
        run.font.size = Pt(font_size)
        if bold:
            run.font.bold = True
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    return para

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with open(DOC_PATH, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    code_para = None

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code block
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1.0)
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                run = para.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x17, 0x20, 0x2A)
                # Box shading
                pPr = para._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F4F6F7')
                pPr.append(shd)
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        if line.startswith('#### '):
            para = doc.add_paragraph()
            run = para.add_run(line[5:])
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
            para.paragraph_format.space_before = Pt(6)
            i += 1
            continue
        if line.startswith('### '):
            para = doc.add_heading(line[4:], level=3)
            para.paragraph_format.space_before = Pt(10)
            i += 1
            continue
        if line.startswith('## '):
            para = doc.add_heading(line[3:], level=2)
            para.paragraph_format.space_before = Pt(14)
            i += 1
            continue
        if line.startswith('# '):
            para = doc.add_heading(line[2:], level=1)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Metadata lines (bold key: value pattern at top)
        if line.startswith('**') and line.count('**') >= 2 and i < 10:
            add_formatted_para(doc, line, font_size=11)
            i += 1
            continue

        # Table
        if line.startswith('|'):
            i = add_table_from_md(doc, lines, i)
            continue

        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.space_after = Pt(3)
            segments = parse_inline(text)
            for seg_text, bold, code in segments:
                run = para.add_run(seg_text)
                run.font.size = Pt(11)
                if bold:
                    run.font.bold = True
                if code:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\. (.+)', line)
        if m:
            text = m.group(2)
            para = doc.add_paragraph(style='List Number')
            para.paragraph_format.space_after = Pt(3)
            segments = parse_inline(text)
            for seg_text, bold, code in segments:
                run = para.add_run(seg_text)
                run.font.size = Pt(11)
                if bold:
                    run.font.bold = True
                if code:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        add_formatted_para(doc, line)
        i += 1

    # Footer
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Federated Learning with XAI for Chest X-Ray Classification  |  Advanced ML Project  |  MS 2026"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUT_PATH)
    print(f"Document saved: {OUT_PATH}")

if __name__ == "__main__":
    main()
